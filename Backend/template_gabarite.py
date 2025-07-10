from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_header_footer(doc, nome_projeto="Gabarite – Versão Legislação"):
    for section in doc.sections:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = nome_projeto
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = header_paragraph.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 87, 255)
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Gabarite - www.gabarite.com.br | Página "
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def set_styles(doc):
    styles = doc.styles
    if 'TituloGabarite' not in styles:
        style = styles.add_style('TituloGabarite', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Calibri'
        style.font.size = Pt(22)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 87, 255)
    if 'SubtituloGabarite' not in styles:
        sub = styles.add_style('SubtituloGabarite', WD_STYLE_TYPE.PARAGRAPH)
        sub.font.name = 'Calibri'
        sub.font.size = Pt(16)
        sub.font.bold = True
        sub.font.color.rgb = RGBColor(0, 87, 255)
    if 'TextoGabarite' not in styles:
        norm = styles.add_style('TextoGabarite', WD_STYLE_TYPE.PARAGRAPH)
        norm.font.name = 'Calibri'
        norm.font.size = Pt(12)
        norm.font.color.rgb = RGBColor(20, 20, 20)
    if 'QuadroGabarite' not in styles:
        quadro = styles.add_style('QuadroGabarite', WD_STYLE_TYPE.PARAGRAPH)
        quadro.font.name = 'Calibri'
        quadro.font.size = Pt(11)
        quadro.font.color.rgb = RGBColor(40, 40, 40)
        quadro.font.bold = False

def add_sumario(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldCharType')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldCharType')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run('\n')

def add_titulo(doc, texto):
    p = doc.add_paragraph(texto, style='TituloGabarite')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_subtitulo(doc, texto):
    p = doc.add_paragraph(texto, style='SubtituloGabarite')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_paragrafo(doc, texto):
    p = doc.add_paragraph(texto, style='TextoGabarite')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def add_quadro(doc, titulo, linhas):
    n_cols = len(linhas[0])
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, val in enumerate(linhas[0]):
        hdr_cells[i].text = val
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 87, 255)
    for linha in linhas[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(linha):
            row_cells[i].text = val
    p = doc.add_paragraph(titulo, style='QuadroGabarite')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_pag_break(doc):
    doc.add_page_break()
