from docx.shared import RGBColor

def parse_indice_ia(texto_ia):
    import re
    dados = []
    linhas = [l for l in texto_ia.split('\n') if '|' in l]
    if len(linhas) < 3:
        return dados
    for l in linhas[2:]:
        partes = [p.strip() for p in l.split('|')]
        if len(partes) < 5: continue
        try:
            dados.append({
                "topico": partes[1],
                "subtopico": partes[2],
                "qtd": int(partes[3]),
                "porcentagem": float(str(partes[4]).replace('%','').replace(',','.'))
            })
        except:
            continue
    return dados

def add_indice_estatistico(doc, dados):
    n_cols = 4
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Tópico"
    hdr_cells[1].text = "Subtópico"
    hdr_cells[2].text = "Qtd. Questões"
    hdr_cells[3].text = "% de Incidência"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 87, 255)
    for d in dados:
        row_cells = table.add_row().cells
        row_cells[0].text = d["topico"]
        row_cells[1].text = d["subtopico"]
        row_cells[2].text = str(d["qtd"])
        row_cells[3].text = f'{d["porcentagem"]:.1f}%'
        cor = RGBColor(0, 0, 0)
        if d["porcentagem"] >= 15:
            cor = RGBColor(46, 212, 122)  # Verde
        elif d["porcentagem"] >= 8:
            cor = RGBColor(255, 214, 0)   # Laranja
        else:
            cor = RGBColor(255, 88, 88)   # Vermelho
        for i in range(n_cols):
            for run in row_cells[i].paragraphs[0].runs:
                run.font.color.rgb = cor
    doc.add_paragraph(" ")

def parse_special_blocks(text):
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("[GRIFO"):
            end = "[/GRIFO]"
            categoria = "Padrao"
            if "categoria=" in line:
                categoria = line.split("categoria=")[1].split("]")[0]
            content = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(end):
                content.append(lines[i])
                i += 1
            blocks.append({"type": "grifo", "categoria": categoria, "content": "\n".join(content)})
        elif line.startswith("[QUADRO"):
            end = "[/QUADRO]"
            content = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(end):
                content.append(lines[i])
                i += 1
            blocks.append({"type": "quadro", "content": "\n".join(content)})
        elif line.startswith("[QUESTAO"):
            end = "[/QUESTAO]"
            content = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(end):
                content.append(lines[i])
                i += 1
            blocks.append({"type": "questao", "content": "\n".join(content)})
        elif line != "":
            blocks.append({"type": "texto", "content": line})
        i += 1
    return blocks
