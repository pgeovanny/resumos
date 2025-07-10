from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from enum import Enum
import pdfplumber
import tempfile
import os
from docx import Document
import openai

# IMPORTS RELATIVOS PARA PACOTES NO MESMO DIRETÓRIO
from .template_gabarite import (
    set_header_footer, set_styles, add_sumario, add_titulo, add_subtitulo, add_paragrafo, add_quadro, add_pag_break
)
from .utils import (
    parse_indice_ia, add_indice_estatistico, parse_special_blocks
)

app = FastAPI()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "SUA_API_KEY_AQUI")

class Bancas(str, Enum):
    fgv = "FGV"
    cespe = "Cespe"
    fcc = "FCC"
    verbena = "Verbena"
    outra = "Outra"

def extract_text_from_file(file: UploadFile):
    ext = os.path.splitext(file.filename)[1].lower()
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    temp_file.write(file.file.read())
    temp_file.close()
    text = ""
    if ext == ".txt":
        with open(temp_file.name, encoding="utf-8") as f:
            text = f.read()
    elif ext in [".docx", ".doc"]:
        from docx import Document
        doc = Document(temp_file.name)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == ".csv":
        import pandas as pd
        df = pd.read_csv(temp_file.name, dtype=str)
        text = "\n".join([" | ".join(row) for row in df.values.tolist()])
    else:
        text = ""
    os.unlink(temp_file.name)
    return text

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    paginas: str = Form(None)
):
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    temp_file.write(await file.read())
    temp_file.close()
    ext = os.path.splitext(file.filename)[1]
    if ext.lower() == ".pdf":
        with pdfplumber.open(temp_file.name) as pdf:
            num_pages = len(pdf.pages)
            if paginas:
                page_idxs = []
                for bloco in paginas.split(','):
                    bloco = bloco.strip()
                    if '-' in bloco:
                        ini, fim = map(int, bloco.split('-'))
                        page_idxs.extend(list(range(ini-1, fim)))
                    else:
                        page_idxs.append(int(bloco)-1)
                page_idxs = [i for i in page_idxs if 0 <= i < num_pages]
                text = "\n".join([pdf.pages[i].extract_text() or "" for i in page_idxs])
            else:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
    elif ext.lower() in [".docx", ".doc"]:
        from docx import Document
        doc = Document(temp_file.name)
        text = "\n".join([p.text for p in doc.paragraphs])
    else:
        os.unlink(temp_file.name)
        return JSONResponse({"error": "Formato não suportado"}, status_code=400)
    os.unlink(temp_file.name)
    return {"text": text}

@app.post("/process")
async def process_text(
    text: str = Form(...),
    command: str = Form(...),
    estilo_linguagem: str = Form(None),
    banca: Bancas = Form(...),
    questoes_texto: str = Form(None),
    questoes_file: UploadFile = File(None),
    cargo: str = Form(None),
    ano: str = Form(None)
):
    openai.api_key = OPENAI_API_KEY
    questoes = questoes_texto or ""
    if questoes_file is not None:
        questoes += "\n" + extract_text_from_file(questoes_file)
    prompt = f"""
Você é um especialista em concursos.
Analise as questões a seguir da banca {banca} ({cargo or '[não informado]'}, {ano or '[não informado]'}).
Monte o seguinte quadro ANTES do conteúdo principal:

| Tópico | Subtópico | Qtd. Questões | % de Incidência |
| --- | --- | --- | --- |
| ... (um por linha, ordem decrescente) |

Depois do quadro, produza o material conforme instruções (esquematizado, incluindo as questões reais nos tópicos corretos, sempre explicando como o conteúdo cobre a resposta).
No final, gere 3 questões inéditas no padrão da banca.
Conteúdo base:
{text}
QUESTÕES DA BANCA:
{questoes}

Comando extra: {command}
Linguagem desejada: {estilo_linguagem}
"""
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=4096
    )
    processed = response.choices[0].message.content
    return {"processed_text": processed}

@app.post("/preview")
async def preview_word(processed_text: str = Form(...)):
    doc = Document()
    set_header_footer(doc)
    set_styles(doc)
    add_titulo(doc, "Gabarite – Versão Legislação")
    add_sumario(doc)
    indice_dados = parse_indice_ia(processed_text)
    if indice_dados:
        add_titulo(doc, "Índice Estatístico de Cobrança da Banca")
        add_indice_estatistico(doc, indice_dados)
        add_pag_break(doc)
    blocks = parse_special_blocks(processed_text)
    for block in blocks:
        if block["type"] == "texto":
            add_paragrafo(doc, block["content"])
        elif block["type"] == "grifo":
            p = doc.add_paragraph()
            run = p.add_run(block["content"])
            cor = block.get("categoria", "Padrao").lower()
            if cor == "permissão" or cor == "permissao":
                run.font.highlight_color = 4
            elif cor == "exceção" or cor == "negativa" or cor == "excecao":
                run.font.highlight_color = 6
            elif cor == "atenção" or cor == "atencao" or cor == "prazo":
                run.font.highlight_color = 7
            else:
                run.font.highlight_color = 3
            run.bold = True
            p.add_run(f"  [{block.get('categoria','')}]").italic = True
        elif block["type"] == "quadro":
            linhas = [l for l in block["content"].split("\n") if "|" in l]
            linhas = [list(map(str.strip, l.strip().strip("|").split("|"))) for l in linhas]
            if linhas:
                add_quadro(doc, "Quadro/Esquema", linhas)
        elif block["type"] == "questao":
            add_subtitulo(doc, "Questão (Padrão Banca)")
            for line in block["content"].split("\n"):
                add_paragrafo(doc, line)
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    temp_doc.close()
    return FileResponse(temp_doc.name, filename="gabarite_legislacao_preview.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.post("/generate_word")
async def generate_word(processed_text: str = Form(...)):
    doc = Document()
    set_header_footer(doc)
    set_styles(doc)
    add_titulo(doc, "Gabarite – Versão Legislação")
    add_sumario(doc)
    indice_dados = parse_indice_ia(processed_text)
    if indice_dados:
        add_titulo(doc, "Índice Estatístico de Cobrança da Banca")
        add_indice_estatistico(doc, indice_dados)
        add_pag_break(doc)
    blocks = parse_special_blocks(processed_text)
    for block in blocks:
        if block["type"] == "texto":
            add_paragrafo(doc, block["content"])
        elif block["type"] == "grifo":
            p = doc.add_paragraph()
            run = p.add_run(block["content"])
            cor = block.get("categoria", "Padrao").lower()
            if cor == "permissão" or cor == "permissao":
                run.font.highlight_color = 4
            elif cor == "exceção" or cor == "negativa" or cor == "excecao":
                run.font.highlight_color = 6
            elif cor == "atenção" or cor == "atencao" or cor == "prazo":
                run.font.highlight_color = 7
            else:
                run.font.highlight_color = 3
            run.bold = True
            p.add_run(f"  [{block.get('categoria','')}]").italic = True
        elif block["type"] == "quadro":
            linhas = [l for l in block["content"].split("\n") if "|" in l]
            linhas = [list(map(str.strip, l.strip().strip("|").split("|"))) for l in linhas]
            if linhas:
                add_quadro(doc, "Quadro/Esquema", linhas)
        elif block["type"] == "questao":
            add_subtitulo(doc, "Questão (Padrão Banca)")
            for line in block["content"].split("\n"):
                add_paragrafo(doc, line)
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    temp_doc.close()
    return FileResponse(temp_doc.name, filename="gabarite_legislacao.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
