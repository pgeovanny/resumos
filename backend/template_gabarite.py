from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_header_footer(doc: Document):
    # Adiciona cabeçalho e rodapé simples, editável no Word
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Gabarite – Legislação Facilitada"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Página "
    p.add_run().field = 'PAGE'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_styles(doc: Document):
    # Define estilos básicos do documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(34, 34, 34)

def add_titulo(doc: Document, titulo):
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subtitulo(doc: Document, subtitulo):
    p = doc.add_paragraph()
    run = p.add_run(subtitulo)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_sumario(doc: Document):
    # Cria uma página para sumário (editável pelo usuário depois)
    doc.add_paragraph("SUMÁRIO (Edite manualmente no Word)")
    doc.add_paragraph(" " * 5)

def add_paragrafo(doc: Document, texto):
    doc.add_paragraph(texto)

def add_pag_break(doc: Document):
    doc.add_page_break()

def add_quadro(doc: Document, titulo, linhas):
    # linhas: lista de listas
    table = doc.add_table(rows=1, cols=len(linhas[0]))
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    for i, item in enumerate(linhas[0]):
        hdr_cells[i].text = item
    for linha in linhas[1:]:
        row_cells = table.add_row().cells
        for i, item in enumerate(linha):
            row_cells[i].text = item
    if titulo:
        add_subtitulo(doc, titulo)
